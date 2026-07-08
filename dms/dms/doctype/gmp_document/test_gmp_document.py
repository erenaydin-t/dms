# Copyright (c) 2026, ErenAydin - GMP DMS Module
# License: MIT
"""Unit tests for the GMP Document controller.

Coverage:
    - Autoname rules (per-pair increment, amendment suffix, missing abbr)
    - Amendment lifecycle (version bump + field reset)
    - Lifecycle date calculation (expiry, next revision)
    - SHA-256 integrity helper
    - Watermark resolver state machine
    - Server-side mandatory enforcement of `reason_for_change`

Integration tests that require LibreOffice and real .docx fixtures
(template rendering, base-PDF generation, watermark overlay) are marked
as `skipUnless` so the unit suite remains hermetic.
"""

import os
import shutil
import tempfile
import unittest
from unittest.mock import MagicMock

import frappe
from frappe.tests.utils import FrappeTestCase
from frappe.utils import cint

from dms.dms.doctype.gmp_document.gmp_document import (
    TEMPLATE_FIELDS,
    TEMPLATE_FIELD_KEYS,
    VALIDITY_YEARS_MAP,
    _apply_watermark,
    _compute_sha256,
    _resolve_watermark_text,
    get_template_field_catalog,
)


TEST_DEPT = "GMP-Test-QA Department"
TEST_DEPT_ABBR = "QA"
SOFFICE_AVAILABLE = bool(shutil.which("soffice") or shutil.which("libreoffice"))

# Stop Frappe's test-record generator from recursing into GMP Document's link
# dependencies that reach ERPNext's Company doctype. Importing ERPNext's company
# test module runs erpnext.tests.utils at import time (BootStrapTestData -> Item
# opening stock -> Stock Entry submit), which fails on a version-mismatched
# bench. Company is reached two ways: directly via Department/Employee, and
# transitively via User -> Email Account -> Company. None of these test records
# are needed here — every test builds the Department/Employee/User fixtures it
# uses explicitly (or uses Administrator).
IGNORE_TEST_RECORD_DEPENDENCIES = ["Employee", "Department", "Company", "User"]


def _ensure_test_department():
    """Create a Department fixture with a known custom_abbr."""
    if not frappe.db.exists("Department", TEST_DEPT):
        dept = frappe.new_doc("Department")
        dept.department_name = TEST_DEPT
        dept.is_group = 0
        dept.flags.ignore_mandatory = True
        dept.flags.ignore_permissions = True
        dept.insert(ignore_permissions=True)
    frappe.db.set_value("Department", TEST_DEPT, "custom_abbr", TEST_DEPT_ABBR)
    frappe.db.commit()


def _ensure_test_document_types():
    """document_type is a Link to GMP Document Type, so the codes used by the
    tests must exist as master records (normally seeded by install.after_migrate;
    asserted here so the suite is hermetic on a fresh site)."""
    for type_name, code in (("SOP", "SOP"), ("Work Instruction", "WI")):
        if not frappe.db.exists("GMP Document Type", code):
            frappe.get_doc({
                "doctype": "GMP Document Type",
                "code": code,
                "type_name": type_name,
            }).insert(ignore_permissions=True)
    frappe.db.commit()


TEST_WORD_TEMPLATE = "GMP-Test-Template"


def _ensure_test_word_template():
    """word_template is now mandatory on GMP Document, so a template must exist.
    A template is file-less (title + mappings); one native mapping is enough to
    exercise the alias path. Returns the template name."""
    if not frappe.db.exists("GMP Word Template", TEST_WORD_TEMPLATE):
        frappe.get_doc({
            "doctype": "GMP Word Template",
            "template_title": TEST_WORD_TEMPLATE,
            "field_mappings": [
                {"custom_tag": "my_title", "system_field": "document_name_en"},
            ],
        }).insert(ignore_permissions=True)
    frappe.db.commit()
    return TEST_WORD_TEMPLATE


# Minimal valid PNG (1x1) — enough for the signature validation, which checks
# the file exists and is a PNG/JPG/JPEG, not that it is a real image.
_SIG_PNG = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\xf8\x0f\x00"
    b"\x01\x01\x01\x00\x18\xdd\x8d\xb0\x00\x00\x00\x00IEND\xaeB`\x82"
)


def _ensure_signature(user, dept):
    """Give `user` a linked Employee with a valid PNG signature so the
    Reviewer/QA signature validation (GMPDocument._validate_signatures) passes
    for documents that assign this user as reviewer/qa_approver."""
    emp = frappe.db.get_value("Employee", {"user_id": user}, "name")
    if not emp:
        e = frappe.new_doc("Employee")
        e.name = f"GMP-SIG-EMP-{frappe.generate_hash(length=8)}"
        e.first_name = user.split("@")[0] if "@" in user else user
        e.employee_name = e.first_name
        e.user_id = user
        e.department = dept
        e.status = "Active"
        e.flags.ignore_mandatory = True
        e.db_insert()
        emp = e.name
    if not frappe.db.get_value("Employee", emp, "custom_signature_image"):
        f = frappe.get_doc(
            {
                "doctype": "File",
                "file_name": f"sig-{frappe.generate_hash(length=6)}.png",
                "is_private": 1,
                "content": _SIG_PNG,
            }
        ).insert(ignore_permissions=True)
        frappe.db.set_value("Employee", emp, "custom_signature_image", f.file_url)
    frappe.db.commit()


def _purge_test_documents():
    for name in frappe.get_all(
        "GMP Document",
        filters=[["document_name_en", "like", "GMP-Test-%"]],
        pluck="name",
    ):
        try:
            doc = frappe.get_doc("GMP Document", name)
            if doc.docstatus == 1:
                doc.flags.ignore_permissions = True
                doc.cancel()
            # The audit-retention guard (on_trash) blocks deleting cancelled
            # revisions; downgrade the status so test fixtures can be purged.
            if doc.workflow_status == "Revision Cancelled":
                doc.db_set("workflow_status", "Draft", update_modified=False)
            frappe.delete_doc("GMP Document", name, ignore_permissions=True, force=True)
        except Exception:
            pass
    frappe.db.commit()


class TestGMPDocument(FrappeTestCase):
    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        _ensure_test_department()
        _ensure_test_document_types()
        _ensure_test_word_template()
        # reviewer/qa_approver default to Administrator in _build_doc; the new
        # signature validation requires them to have a signature image.
        _ensure_signature("Administrator", TEST_DEPT)
        _purge_test_documents()

    @classmethod
    def tearDownClass(cls):
        _purge_test_documents()
        super().tearDownClass()

    def tearDown(self):
        _purge_test_documents()

    # ------------------------------------------------------------------ #
    #  Helpers                                                           #
    # ------------------------------------------------------------------ #

    def _dummy_attachment(self, en):
        """A minimal .docx-named private File with arbitrary bytes. Valid for
        save/hash/rename; only actual rendering (on_submit) needs a real .docx,
        which the LibreOffice end-to-end tests supply via _make_docx_file().

        Content must be unique per file: Frappe deduplicates File records by
        content hash, so identical bytes would collapse distinct uploads onto a
        single File/URL and break tests that depend on a genuinely fresh file
        (e.g. the amend "keeps freshly uploaded file" case)."""
        nonce = frappe.generate_hash(length=16)
        fname = f"{en}-{frappe.generate_hash(length=6)}.docx"
        return frappe.get_doc({
            "doctype": "File",
            "file_name": fname,
            "is_private": 1,
            "content": b"PK\x03\x04 dummy docx for tests " + nonce.encode(),
        }).insert(ignore_permissions=True)

    def _build_doc(self, **overrides):
        defaults = {
            "doctype": "GMP Document",
            "document_name_fa": "تست",
            "document_name_en": "GMP-Test-Default",
            "document_type": "SOP",
            "department": TEST_DEPT,
            "gmp_impact": "Major",
            "validity_period": "3 Years",
            "version_number": 1,
            # reviewer and qa_approver are mandatory; default to Administrator
            # so tests that don't exercise the workflow can still insert.
            "reviewer": "Administrator",
            "qa_approver": "Administrator",
        }
        defaults.update(overrides)
        doc = frappe.new_doc("GMP Document")
        doc.update(defaults)
        # word_template and attachment_file are mandatory; satisfy them unless a
        # test overrides them explicitly.
        if "word_template" not in overrides:
            doc.word_template = _ensure_test_word_template()
        if "attachment_file" not in overrides:
            doc.attachment_file = self._dummy_attachment(defaults["document_name_en"]).file_url
        return doc

    def _approve_via_workflow(self, doc):
        """Drive a draft through the native Workflow to Approved (submitted).

        Transitions go through Frappe's apply_workflow (the same path the
        "Actions" menu uses). Tests run as Administrator, which satisfies every
        transition role and the per-actor conditions' Administrator escape."""
        from frappe.model.workflow import apply_workflow

        apply_workflow(doc, "Submit for Review")
        apply_workflow(doc, "Approve as Reviewer")
        apply_workflow(doc, "Approve as QA")
        doc.reload()
        return doc

    # ------------------------------------------------------------------ #
    #  Autoname                                                          #
    # ------------------------------------------------------------------ #

    def test_autoname_first_doc_in_pair(self):
        doc = self._build_doc(document_name_en="GMP-Test-First")
        doc.insert(ignore_permissions=True)
        self.assertEqual(doc.name, f"{TEST_DEPT_ABBR}-SOP-0001-1")

    def test_autoname_increments_per_type_dept_pair(self):
        d1 = self._build_doc(document_name_en="GMP-Test-Inc-1")
        d1.insert(ignore_permissions=True)
        d2 = self._build_doc(document_name_en="GMP-Test-Inc-2")
        d2.insert(ignore_permissions=True)

        self.assertEqual(d1.name, f"{TEST_DEPT_ABBR}-SOP-0001-1")
        self.assertEqual(d2.name, f"{TEST_DEPT_ABBR}-SOP-0002-1")

    def test_autoname_separate_increments_for_different_types(self):
        sop = self._build_doc(document_type="SOP", document_name_en="GMP-Test-SOP")
        sop.insert(ignore_permissions=True)
        wi = self._build_doc(document_type="WI", document_name_en="GMP-Test-WI")
        wi.insert(ignore_permissions=True)

        self.assertEqual(sop.name, f"{TEST_DEPT_ABBR}-SOP-0001-1")
        self.assertEqual(wi.name, f"{TEST_DEPT_ABBR}-WI-0001-1")

    def test_autoname_amended_doc_keeps_base_id(self):
        original = self._build_doc(document_name_en="GMP-Test-Amend-Origin")
        original.insert(ignore_permissions=True)

        amended = self._build_doc(document_name_en="GMP-Test-Amend-V1")
        amended.amended_from = original.name
        amended.reason_for_change = "Procedural correction per CAPA-2026-001"
        amended.before_insert()
        amended.autoname()

        base = original.name.rsplit("-", 1)[0]
        self.assertEqual(amended.name, f"{base}-2")
        self.assertEqual(amended.version_number, 2)

    def test_autoname_throws_when_dept_missing_abbr(self):
        no_abbr_dept = "GMP-Test-NoAbbr Department"
        if not frappe.db.exists("Department", no_abbr_dept):
            d = frappe.new_doc("Department")
            d.department_name = no_abbr_dept
            d.flags.ignore_mandatory = True
            d.insert(ignore_permissions=True)
        frappe.db.set_value("Department", no_abbr_dept, "custom_abbr", "")
        frappe.db.commit()

        doc = self._build_doc(department=no_abbr_dept, document_name_en="GMP-Test-NoAbbr")
        with self.assertRaises(frappe.ValidationError):
            doc.insert(ignore_permissions=True)

    # ------------------------------------------------------------------ #
    #  Amendment lifecycle                                               #
    # ------------------------------------------------------------------ #

    def test_amendment_clears_inherited_file_and_bumps_version(self):
        original = self._build_doc(document_name_en="GMP-Test-Reset-Origin")
        original.insert(ignore_permissions=True)

        amended = self._build_doc(
            document_name_en="GMP-Test-Reset-V1",
            # Inherited from the predecessor (as the amend copy would carry):
            attachment_file=original.attachment_file,
        )
        amended.amended_from = original.name
        amended.file_integrity_hash = "deadbeef" * 8
        amended.effective_date = "2026-04-01"
        amended.expiry_date = "2029-04-01"
        amended.next_revision_date = "2029-03-01"

        amended.before_insert()

        self.assertEqual(amended.version_number, 2)
        # An inherited file must be cleared so the user re-uploads a fresh one.
        self.assertIsNone(amended.attachment_file)
        self.assertIsNone(amended.file_integrity_hash)
        self.assertIsNone(amended.effective_date)
        self.assertIsNone(amended.expiry_date)
        self.assertIsNone(amended.next_revision_date)

    def test_amendment_keeps_freshly_uploaded_file(self):
        original = self._build_doc(document_name_en="GMP-Test-Fresh-Origin")
        original.insert(ignore_permissions=True)

        fresh_url = self._dummy_attachment("GMP-Test-Fresh-New").file_url
        amended = self._build_doc(
            document_name_en="GMP-Test-Fresh-V1",
            attachment_file=fresh_url,  # user uploaded a new file for the revision
        )
        amended.amended_from = original.name

        amended.before_insert()

        self.assertEqual(amended.version_number, 2)
        self.assertEqual(amended.attachment_file, fresh_url)

    # ------------------------------------------------------------------ #
    #  Non-destructive revision lifecycle (create_revision)              #
    # ------------------------------------------------------------------ #

    def _force_approved(self, doc):
        """Simulate a QA-approved effective version without running the render
        pipeline (mirrors the docstatus-forcing pattern of the amend tests)."""
        frappe.db.set_value(
            "GMP Document",
            doc.name,
            {"docstatus": 1, "workflow_status": "Approved", "is_active": 1},
            update_modified=False,
        )
        doc.reload()
        return doc

    def test_create_revision_keeps_predecessor_effective(self):
        from dms.dms.doctype.gmp_document.gmp_document import create_revision

        original = self._build_doc(document_name_en="GMP-Test-Rev-Origin")
        original.insert(ignore_permissions=True)
        self._force_approved(original)

        rev_name = create_revision(original.name, "Update per CAPA-2026-002")
        rev = frappe.get_doc("GMP Document", rev_name)

        # The draft revision is a separate, linked record with a fresh cycle.
        self.assertEqual(rev.revision_of, original.name)
        self.assertEqual(rev.docstatus, 0)
        self.assertEqual(rev.workflow_status, "Draft")
        self.assertEqual(rev.version_number, 2)
        self.assertEqual(rev.name, f"{original.name.rsplit('-', 1)[0]}-2")
        # The predecessor's controlled file is never carried over.
        self.assertFalse(rev.attachment_file)
        self.assertFalse(rev.effective_date)

        # The predecessor is untouched: still Approved, submitted and active.
        original.reload()
        self.assertEqual(original.docstatus, 1)
        self.assertEqual(original.workflow_status, "Approved")
        self.assertEqual(cint(original.is_active), 1)
        self.assertEqual(original.version_number, 1)

    def test_create_revision_requires_approved_active_source(self):
        from dms.dms.doctype.gmp_document.gmp_document import create_revision

        draft = self._build_doc(document_name_en="GMP-Test-Rev-Draft")
        draft.insert(ignore_permissions=True)
        with self.assertRaises(frappe.ValidationError):
            create_revision(draft.name, "reason")

        retired = self._build_doc(document_name_en="GMP-Test-Rev-Retired")
        retired.insert(ignore_permissions=True)
        self._force_approved(retired)
        frappe.db.set_value("GMP Document", retired.name, "is_active", 0, update_modified=False)
        with self.assertRaises(frappe.ValidationError):
            create_revision(retired.name, "reason")

    def test_create_revision_requires_reason(self):
        from dms.dms.doctype.gmp_document.gmp_document import create_revision

        original = self._build_doc(document_name_en="GMP-Test-Rev-NoReason")
        original.insert(ignore_permissions=True)
        self._force_approved(original)

        with self.assertRaises(frappe.ValidationError):
            create_revision(original.name, "   ")

    def test_create_revision_blocks_parallel_open_revision(self):
        from dms.dms.doctype.gmp_document.gmp_document import create_revision

        original = self._build_doc(document_name_en="GMP-Test-Rev-Parallel")
        original.insert(ignore_permissions=True)
        self._force_approved(original)

        create_revision(original.name, "first attempt")
        with self.assertRaises(frappe.ValidationError):
            create_revision(original.name, "second attempt")

    def test_cancelled_revision_frees_slot_and_next_name_skips_it(self):
        from dms.dms.doctype.gmp_document.gmp_document import create_revision

        original = self._build_doc(document_name_en="GMP-Test-Rev-Cancelled")
        original.insert(ignore_permissions=True)
        self._force_approved(original)
        base = original.name.rsplit("-", 1)[0]

        first = create_revision(original.name, "abandoned attempt")
        self.assertEqual(first, f"{base}-2")
        # Simulate the 'Cancel Revision' workflow transition.
        frappe.db.set_value(
            "GMP Document",
            first,
            {"workflow_status": "Revision Cancelled", "is_active": 0},
            update_modified=False,
        )

        # The slot is free again, and the retained cancelled record keeps its
        # name — the next revision takes the NEXT segment, never colliding.
        second = create_revision(original.name, "second attempt")
        self.assertEqual(second, f"{base}-3")
        self.assertEqual(frappe.db.get_value("GMP Document", second, "version_number"), 2)

        # Predecessor still effective throughout.
        original.reload()
        self.assertEqual(original.workflow_status, "Approved")
        self.assertEqual(cint(original.is_active), 1)

    def test_revision_approval_obsoletes_predecessor(self):
        from dms.dms.doctype.gmp_document.gmp_document import create_revision

        original = self._build_doc(document_name_en="GMP-Test-Rev-Supersede")
        original.insert(ignore_permissions=True)
        self._force_approved(original)

        rev = frappe.get_doc("GMP Document", create_revision(original.name, "supersede test"))
        # Exercise the on_submit hand-over directly (the full submit path needs
        # LibreOffice and is covered by the e2e suite).
        rev._obsolete_superseded_version()

        original.reload()
        self.assertEqual(original.workflow_status, "Obsolete")
        self.assertEqual(cint(original.is_active), 0)
        # Obsolete via supersession is NOT a cancel: the record stays submitted.
        self.assertEqual(original.docstatus, 1)

    def test_cancelled_revision_cannot_be_deleted(self):
        from dms.dms.doctype.gmp_document.gmp_document import create_revision

        original = self._build_doc(document_name_en="GMP-Test-Rev-Retain")
        original.insert(ignore_permissions=True)
        self._force_approved(original)

        rev_name = create_revision(original.name, "to be cancelled")
        frappe.db.set_value(
            "GMP Document", rev_name, "workflow_status", "Revision Cancelled", update_modified=False
        )

        with self.assertRaises(frappe.ValidationError):
            frappe.delete_doc("GMP Document", rev_name, ignore_permissions=True)

    # ------------------------------------------------------------------ #
    #  Effective-date scheduling (posting-date semantics)                 #
    # ------------------------------------------------------------------ #

    def test_effective_date_is_system_controlled_without_checkbox(self):
        doc = self._build_doc(document_name_en="GMP-Test-Eff-SysCtl")
        doc.effective_date = "2030-01-01"  # manual value without the checkbox
        doc.insert(ignore_permissions=True)
        # Silently normalized away, like ERPNext's posting date.
        self.assertFalse(doc.effective_date)
        self.assertFalse(doc.expiry_date)

    def test_effective_date_manual_with_checkbox_survives(self):
        doc = self._build_doc(document_name_en="GMP-Test-Eff-Manual")
        doc.set_effective_date = 1
        doc.effective_date = "2030-01-01"
        doc.insert(ignore_permissions=True)  # Administrator passes the role gate
        self.assertEqual(str(doc.effective_date), "2030-01-01")
        # Lifecycle dates derive from the scheduled date (3 Years validity).
        self.assertEqual(str(doc.expiry_date), "2033-01-01")

    def test_submit_requires_date_when_manual_mode_enabled(self):
        doc = self._build_doc(document_name_en="GMP-Test-Eff-NoDate")
        doc.set_effective_date = 1
        doc.insert(ignore_permissions=True)
        doc.workflow_status = "Approved"  # satisfy the before_submit guard
        with self.assertRaises(frappe.ValidationError):
            doc.submit()

    def test_activation_sweep_promotes_due_pending_document(self):
        from dms.dms.doctype.gmp_document.gmp_document import activate_effective_documents

        original = self._build_doc(document_name_en="GMP-Test-Eff-Sweep-Origin")
        original.insert(ignore_permissions=True)
        self._force_approved(original)

        # A pending revision whose effective date has arrived.
        from dms.dms.doctype.gmp_document.gmp_document import create_revision
        rev_name = create_revision(original.name, "scheduled rollout")
        frappe.db.set_value(
            "GMP Document",
            rev_name,
            {
                "docstatus": 1,
                "workflow_status": "Approved",
                "is_active": 0,
                "effective_date": frappe.utils.add_days(frappe.utils.today(), -1),
            },
            update_modified=False,
        )

        activate_effective_documents()

        rev = frappe.get_doc("GMP Document", rev_name)
        self.assertEqual(cint(rev.is_active), 1)
        original.reload()
        self.assertEqual(original.workflow_status, "Obsolete")
        self.assertEqual(cint(original.is_active), 0)
        self.assertEqual(original.docstatus, 1)

    def test_activation_sweep_skips_future_and_obsolete(self):
        from dms.dms.doctype.gmp_document.gmp_document import activate_effective_documents

        future = self._build_doc(document_name_en="GMP-Test-Eff-Future")
        future.insert(ignore_permissions=True)
        frappe.db.set_value(
            "GMP Document",
            future.name,
            {
                "docstatus": 1,
                "workflow_status": "Approved",
                "is_active": 0,
                "effective_date": frappe.utils.add_days(frappe.utils.today(), 10),
            },
            update_modified=False,
        )

        obsolete = self._build_doc(document_name_en="GMP-Test-Eff-Obsolete")
        obsolete.insert(ignore_permissions=True)
        frappe.db.set_value(
            "GMP Document",
            obsolete.name,
            {
                "docstatus": 1,
                "workflow_status": "Obsolete",
                "is_active": 0,
                "effective_date": frappe.utils.add_days(frappe.utils.today(), -30),
            },
            update_modified=False,
        )

        activate_effective_documents()

        self.assertEqual(
            cint(frappe.db.get_value("GMP Document", future.name, "is_active")), 0
        )
        self.assertEqual(
            frappe.db.get_value("GMP Document", obsolete.name, "workflow_status"), "Obsolete"
        )
        self.assertEqual(
            cint(frappe.db.get_value("GMP Document", obsolete.name, "is_active")), 0
        )

    # ------------------------------------------------------------------ #
    #  Approver stamping (drives the approver's PDF signature)            #
    # ------------------------------------------------------------------ #

    def test_stamp_approver_fills_when_missing(self):
        # The approver's signature is resolved from approved_by during render;
        # on_submit must guarantee it is stamped before that, independent of the
        # on_update workflow side-effect.
        doc = self._build_doc(document_name_en="GMP-Test-Approver")
        doc.insert(ignore_permissions=True)
        self.assertFalse(doc.approved_by)

        doc._stamp_approver()

        self.assertEqual(doc.approved_by, frappe.session.user)
        self.assertIsNotNone(doc.approved_on)

    def test_stamp_approver_is_idempotent(self):
        doc = self._build_doc(document_name_en="GMP-Test-Approver2")
        doc.insert(ignore_permissions=True)
        doc.db_set("approved_by", "Administrator", update_modified=False)

        doc._stamp_approver()

        self.assertEqual(doc.approved_by, "Administrator")

    def test_amendment_requires_reason_for_change(self):
        original = self._build_doc(document_name_en="GMP-Test-ReasonReq-Origin")
        original.insert(ignore_permissions=True)

        amended = self._build_doc(document_name_en="GMP-Test-ReasonReq-V1")
        amended.amended_from = original.name
        amended.reason_for_change = ""  # missing

        with self.assertRaises(frappe.ValidationError):
            amended.insert(ignore_permissions=True)

    # ------------------------------------------------------------------ #
    #  Reviewer / QA signature validation                                 #
    # ------------------------------------------------------------------ #

    def _user_without_signature(self, email):
        if not frappe.db.exists("User", email):
            u = frappe.new_doc("User")
            u.email = email
            u.first_name = email.split("@")[0]
            u.user_type = "System User"
            u.send_welcome_email = 0
            u.insert(ignore_permissions=True)
        # make sure any linked Employee has no signature
        emp = frappe.db.get_value("Employee", {"user_id": email}, "name")
        if emp:
            frappe.db.set_value("Employee", emp, "custom_signature_image", "")
        frappe.db.commit()
        return email

    def test_reviewer_without_signature_blocks_save(self):
        bad = self._user_without_signature("gmp-test-nosig-reviewer@example.com")
        doc = self._build_doc(document_name_en="GMP-Test-NoSigReviewer", reviewer=bad)
        with self.assertRaises(frappe.ValidationError):
            doc.insert(ignore_permissions=True)

    def test_qa_approver_without_signature_blocks_save(self):
        bad = self._user_without_signature("gmp-test-nosig-qa@example.com")
        doc = self._build_doc(document_name_en="GMP-Test-NoSigQA", qa_approver=bad)
        with self.assertRaises(frappe.ValidationError):
            doc.insert(ignore_permissions=True)

    def test_reviewer_and_qa_with_signature_save_ok(self):
        # Administrator has a signature (ensured in setUpClass) -> save succeeds.
        doc = self._build_doc(document_name_en="GMP-Test-SigOK")
        doc.insert(ignore_permissions=True)
        self.assertTrue(doc.name)

    # ------------------------------------------------------------------ #
    #  Word template field mapping                                        #
    # ------------------------------------------------------------------ #

    def test_template_context_includes_native_keys(self):
        doc = self._build_doc(document_name_en="GMP-Test-Ctx")
        ctx = doc._build_template_context()
        self.assertEqual(ctx["document_name_en"], "GMP-Test-Ctx")
        # Native keys remain available with no mappings (backward compatible).
        self.assertIn("version_number", ctx)

    def test_custom_tag_alias_mirrors_system_field(self):
        doc = self._build_doc(document_name_en="GMP-Test-Alias")
        mappings = [
            frappe._dict(custom_tag="my_title", system_field="document_name_en"),
            frappe._dict(custom_tag="rev_no", system_field="version_number"),
        ]
        ctx = doc._build_template_context(field_mappings=mappings)
        self.assertEqual(ctx["my_title"], ctx["document_name_en"])
        self.assertEqual(ctx["rev_no"], ctx["version_number"])

    def test_alias_to_unknown_system_field_is_ignored(self):
        doc = self._build_doc()
        mappings = [frappe._dict(custom_tag="ghost", system_field="not_a_field")]
        ctx = doc._build_template_context(field_mappings=mappings)
        self.assertNotIn("ghost", ctx)

    def test_field_catalog_matches_context_keys(self):
        # Every catalog key must be a real context key, or the mapping UI would
        # offer fields that render blank.
        ctx = self._build_doc()._build_template_context()
        for key in TEMPLATE_FIELD_KEYS:
            self.assertIn(key, ctx, f"catalog key '{key}' missing from template context")

    def test_get_template_field_catalog_shape(self):
        catalog = get_template_field_catalog()
        self.assertEqual(len(catalog), len(TEMPLATE_FIELDS))
        self.assertEqual(catalog[0].keys(), {"value", "label"})

    # ------------------------------------------------------------------ #
    #  Lifecycle date calculation                                        #
    # ------------------------------------------------------------------ #

    def test_expiry_and_revision_date_calculation(self):
        doc = self._build_doc(validity_period="3 Years")
        doc.effective_date = "2026-01-15"
        doc._calculate_lifecycle_dates()

        self.assertEqual(str(doc.expiry_date), "2029-01-15")
        self.assertEqual(str(doc.next_revision_date), "2028-12-15")

    def test_lifecycle_dates_skipped_without_effective_date(self):
        doc = self._build_doc()
        doc.effective_date = None
        doc.expiry_date = None
        doc.next_revision_date = None
        doc._calculate_lifecycle_dates()

        self.assertIsNone(doc.expiry_date)
        self.assertIsNone(doc.next_revision_date)

    def test_validity_period_map_covers_all_options(self):
        # Guards against future Select-option drift.
        self.assertEqual(VALIDITY_YEARS_MAP, {"2 Years": 2, "3 Years": 3, "5 Years": 5})

    # ------------------------------------------------------------------ #
    #  SHA-256 integrity                                                 #
    # ------------------------------------------------------------------ #

    def test_compute_sha256_is_deterministic_and_correct_length(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".bin") as fh:
            fh.write(b"GMP integrity test content")
            path = fh.name
        try:
            digest = _compute_sha256(path)
            self.assertEqual(len(digest), 64)
            self.assertEqual(digest, _compute_sha256(path))
        finally:
            os.unlink(path)

    def test_compute_sha256_changes_with_content(self):
        with tempfile.NamedTemporaryFile(delete=False) as fh:
            fh.write(b"alpha")
            path_a = fh.name
        with tempfile.NamedTemporaryFile(delete=False) as fh:
            fh.write(b"beta")
            path_b = fh.name
        try:
            self.assertNotEqual(_compute_sha256(path_a), _compute_sha256(path_b))
        finally:
            os.unlink(path_a)
            os.unlink(path_b)

    # ------------------------------------------------------------------ #
    #  Watermark state machine                                           #
    # ------------------------------------------------------------------ #

    def test_watermark_controlled_copy(self):
        doc = MagicMock(docstatus=1, is_active=1)
        self.assertEqual(_resolve_watermark_text(doc), "CONTROLLED COPY")

    def test_watermark_obsolete_when_inactive(self):
        doc = MagicMock(docstatus=1, is_active=0)
        self.assertEqual(_resolve_watermark_text(doc), "OBSOLETE")

    def test_watermark_obsolete_overrides_draft_state(self):
        # Cancelled draft (docstatus=2) is also inactive -> OBSOLETE wins.
        doc = MagicMock(docstatus=2, is_active=0)
        self.assertEqual(_resolve_watermark_text(doc), "OBSOLETE")

    def test_watermark_draft_for_unsubmitted_active(self):
        doc = MagicMock(docstatus=0, is_active=1)
        self.assertEqual(_resolve_watermark_text(doc), "DRAFT - NOT FOR USE")

    def test_watermark_not_yet_effective_for_pending(self):
        # QA-approved but future-dated: neither controlled nor obsolete.
        doc = MagicMock(
            docstatus=1,
            is_active=0,
            workflow_status="Approved",
            effective_date=frappe.utils.add_days(frappe.utils.today(), 5),
        )
        self.assertEqual(_resolve_watermark_text(doc), "NOT YET EFFECTIVE")

    def test_watermark_pending_until_sweep_activates(self):
        # The date has arrived but the activation sweep hasn't run yet: the
        # document is still pending — it must NOT read as Obsolete.
        doc = MagicMock(
            docstatus=1,
            is_active=0,
            workflow_status="Approved",
            effective_date=frappe.utils.add_days(frappe.utils.today(), -1),
        )
        self.assertEqual(_resolve_watermark_text(doc), "NOT YET EFFECTIVE")

    # ------------------------------------------------------------------ #
    #  Watermark overlay (integration - requires reportlab + pypdf)      #
    # ------------------------------------------------------------------ #

    def test_apply_watermark_returns_pdf_bytes(self):
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas

        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as fh:
            c = canvas.Canvas(fh.name, pagesize=A4)
            c.drawString(100, 750, "Base content")
            c.save()
            base_path = fh.name

        try:
            output = _apply_watermark(base_path, "CONTROLLED COPY")
            self.assertTrue(output.startswith(b"%PDF"))
            self.assertGreater(len(output), 0)
        finally:
            os.unlink(base_path)

    # ------------------------------------------------------------------ #
    #  End-to-end submit (skipped unless LibreOffice is on PATH)         #
    # ------------------------------------------------------------------ #

    def _make_docx_file(self, fname="test_template.docx"):
        """Persist a minimal valid .docx as a private File, or skip if
        python-docx is unavailable."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            self.skipTest("python-docx not installed; skipping end-to-end test")

        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, fname)
            d = DocxDocument()
            d.add_paragraph("Document: {{ docname }}")
            d.add_paragraph("Version: {{ version_number }}")
            d.save(docx_path)
            with open(docx_path, "rb") as fh:
                return frappe.get_doc({
                    "doctype": "File",
                    "file_name": fname,
                    "is_private": 1,
                    "content": fh.read(),
                }).insert(ignore_permissions=True)

    def _make_submitted_doc(self, en):
        """Create a GMP Document, attach a real .docx, and drive it through the
        native workflow to Approved/submitted. Returns the reloaded doc."""
        file_doc = self._make_docx_file(fname=f"{en}.docx")
        doc = self._build_doc(document_name_en=en)
        doc.attachment_file = file_doc.file_url
        doc.reviewer = "Administrator"
        doc.qa_approver = "Administrator"
        doc.insert(ignore_permissions=True)
        return self._approve_via_workflow(doc)

    @unittest.skipUnless(SOFFICE_AVAILABLE, "LibreOffice (soffice) not found on PATH")
    def test_submit_persists_base_pdf(self):
        doc = self._make_submitted_doc("GMP-Test-E2E")

        self.assertEqual(doc.docstatus, 1)
        self.assertIsNotNone(doc.effective_date)
        self.assertIsNotNone(doc.file_integrity_hash)
        self.assertEqual(len(doc.file_integrity_hash), 64)

        base_pdf = frappe.db.exists(
            "File",
            {
                "attached_to_doctype": "GMP Document",
                "attached_to_name": doc.name,
                "file_name": f"{doc.name}.pdf",
            },
        )
        self.assertTrue(base_pdf, "Base PDF should be persisted on submit")

    # ------------------------------------------------------------------ #
    #  Issue #3 — amendment naming via the real insert path              #
    # ------------------------------------------------------------------ #

    def test_amend_insert_produces_versioned_name(self):
        # The 'Default Naming' amend rule must let our autoname() run via the
        # real insert path so the name becomes …-0001-2 (version segment bumped),
        # not the framework's …-0001-1-1 counter. Hermetic: no submit/LibreOffice
        # needed — the
        # predecessor is forced to a cancelled docstatus so it is amendable.
        from dms.install import _ensure_amend_naming_rule

        _ensure_amend_naming_rule()

        original = self._build_doc(document_name_en="GMP-Test-AmendName")
        original.insert(ignore_permissions=True)
        frappe.db.set_value("GMP Document", original.name, "docstatus", 2)

        amended = frappe.copy_doc(original)
        amended.docstatus = 0  # amend starts a fresh draft; copy_doc may carry the source docstatus
        amended.amended_from = original.name
        amended.reason_for_change = "CAPA-2026-002 procedural fix"
        # attachment_file is no_copy (not carried by the amend copy) and is now
        # mandatory, so the revision must supply its own controlled file.
        amended.attachment_file = self._dummy_attachment("GMP-Test-AmendName-V1").file_url
        amended.insert(ignore_permissions=True)

        base = original.name.rsplit("-", 1)[0]
        self.assertEqual(amended.name, f"{base}-2")
        # The framework's default amend counter would append to the full source
        # name (…-0001-1-1); our autoname must instead bump the version segment.
        self.assertFalse(amended.name.startswith(original.name + "-"))
        self.assertEqual(amended.version_number, 2)

    # ------------------------------------------------------------------ #
    #  Issue #2 — cancel transitions status to Obsolete                  #
    # ------------------------------------------------------------------ #

    @unittest.skipUnless(SOFFICE_AVAILABLE, "LibreOffice (soffice) not found on PATH")
    def test_cancel_sets_obsolete_status(self):
        doc = self._make_submitted_doc("GMP-Test-Cancel")
        doc.cancel()
        doc.reload()

        self.assertEqual(doc.docstatus, 2)
        self.assertEqual(doc.is_active, 0)
        self.assertEqual(doc.workflow_status, "Obsolete")

    # ------------------------------------------------------------------ #
    #  Issue #4 — references repoint to the new version on amend         #
    # ------------------------------------------------------------------ #

    @unittest.skipUnless(SOFFICE_AVAILABLE, "LibreOffice (soffice) not found on PATH")
    def test_amend_repoints_dependent_references(self):
        from dms.install import _ensure_amend_naming_rule
        from frappe.model.workflow import apply_workflow

        _ensure_amend_naming_rule()

        target = self._make_submitted_doc("GMP-Test-RefTarget")

        # A dependent document referencing the target version.
        dependent = self._make_docx_file(fname="GMP-Test-RefDependent.docx")
        dep = self._build_doc(document_name_en="GMP-Test-RefDependent")
        dep.attachment_file = dependent.file_url
        dep.reviewer = "Administrator"
        dep.qa_approver = "Administrator"
        dep.append("references", {"referenced_document": target.name, "reference_type": "References"})
        dep.insert(ignore_permissions=True)
        dep = self._approve_via_workflow(dep)

        # Cancel + amend the target — link validation must not block this.
        target.cancel()
        amended = frappe.copy_doc(target)
        amended.docstatus = 0  # amend starts a fresh draft; copy_doc may carry the source docstatus
        amended.amended_from = target.name
        amended.reason_for_change = "Revised per change control"
        amended.attachment_file = self._make_docx_file(fname="GMP-Test-RefTarget-v1.docx").file_url
        amended.insert(ignore_permissions=True)
        amended = self._approve_via_workflow(amended)

        # The dependent's reference row should now point at the new version.
        repointed = frappe.get_all(
            "GMP Document Reference",
            filters={"parent": dep.name},
            pluck="referenced_document",
        )
        self.assertIn(amended.name, repointed)
        self.assertNotIn(target.name, repointed)

    # ------------------------------------------------------------------ #
    #  Reference tree — dangling reference (deleted target) is handled    #
    # ------------------------------------------------------------------ #

    def test_reference_tree_handles_deleted_target(self):
        """Regression (v1.2.2): a reference whose target has been deleted must
        not crash the reference tree. Before the fix the per-document permission
        check loaded the target via frappe.get_doc and raised DoesNotExistError
        on a dangling reference; now missing targets are skipped and the rest of
        the tree still renders. Hermetic — no LibreOffice/submit needed."""
        from dms.dms.doctype.gmp_document.gmp_document import get_document_reference_tree

        survivor = self._build_doc(document_name_en="GMP-Test-RefTree-Survivor")
        survivor.insert(ignore_permissions=True)
        doomed = self._build_doc(document_name_en="GMP-Test-RefTree-Doomed")
        doomed.insert(ignore_permissions=True)

        root = self._build_doc(document_name_en="GMP-Test-RefTree-Root")
        root.append("references", {"referenced_document": survivor.name, "reference_type": "References"})
        root.append("references", {"referenced_document": doomed.name, "reference_type": "References"})
        root.insert(ignore_permissions=True)

        # Hard-delete the target to leave a dangling reference row on root
        # (force bypasses the back-link guard, as a manual delete would).
        frappe.delete_doc("GMP Document", doomed.name, ignore_permissions=True, force=True)
        self.assertFalse(frappe.db.exists("GMP Document", doomed.name))

        # Must render without raising (previously DoesNotExistError -> 500).
        tree = get_document_reference_tree(root.name)

        self.assertEqual(tree["name"], root.name)
        child_names = [c["name"] for c in tree["children"]]
        # The deleted target is omitted; the surviving reference still renders.
        self.assertNotIn(doomed.name, child_names)
        self.assertIn(survivor.name, child_names)

    def test_reference_tree_missing_root_raises_not_found(self):
        """A non-existent root docname yields a clean DoesNotExistError, not an
        uncaught crash from frappe.get_doc inside the permission check."""
        from dms.dms.doctype.gmp_document.gmp_document import get_document_reference_tree

        with self.assertRaises(frappe.DoesNotExistError):
            get_document_reference_tree("NOPE-SOP-9999-1")

    # ------------------------------------------------------------------ #
    #  Issue #1 — base PDF regenerates when the File record is missing   #
    # ------------------------------------------------------------------ #

    @unittest.skipUnless(SOFFICE_AVAILABLE, "LibreOffice (soffice) not found on PATH")
    def test_download_regenerates_missing_base_pdf(self):
        from dms.dms.doctype.gmp_document.gmp_document import download_watermarked_pdf

        doc = self._make_submitted_doc("GMP-Test-Regen")

        # Delete the persisted base PDF to simulate the Issue #1 failure mode.
        for pdf in frappe.get_all(
            "File",
            filters={
                "attached_to_doctype": "GMP Document",
                "attached_to_name": doc.name,
                "file_name": ["like", "%.pdf"],
            },
            pluck="name",
        ):
            frappe.delete_doc("File", pdf, ignore_permissions=True, force=True)

        # Should regenerate transparently instead of throwing.
        download_watermarked_pdf(doc.name)
        self.assertEqual(frappe.local.response.type, "download")
        self.assertTrue(frappe.local.response.filecontent.startswith(b"%PDF"))
