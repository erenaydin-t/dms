# Copyright (c) 2026, ErenAydin - GMP DMS Module
# License: MIT
"""End-to-end PDF / template / multi-version / role-access validation.

These tests exercise the *real* pipeline: a user uploads a .docx, the document
is driven through the native workflow to Approved (which renders the Word
template and converts it to PDF via LibreOffice), and the generated PDF bytes
are then extracted and compared. They require `soffice` on PATH and so are
skipped on a bench without LibreOffice.

Covered:
    8A  Template differentiation (no cross-contamination)
    8B  Multi-version content correctness (v0..v4)
    8C  Independent documents (unique PDFs, isolated references)
    9   Role-based access to a real approved document + its PDF
    10  Adversarial regressions (direct-submit bypass, cross-dept PDF, stale
        version content, signature insertion)
"""

import shutil
import unittest

import frappe
from frappe.tests.utils import FrappeTestCase
from frappe.model.workflow import apply_workflow

from dms.dms.doctype.gmp_document.gmp_document import (
    download_watermarked_pdf,
    download_word_document,
    get_document_reference_tree,
)

SOFFICE_AVAILABLE = bool(shutil.which("soffice") or shutil.which("libreoffice"))

DEPT_QA = "GMP-E2E-QA Department"
DEPT_PROD = "GMP-E2E-PROD Department"
TPL_ALPHA = "GMP-E2E-Template-Alpha"
TPL_BETA = "GMP-E2E-Template-Beta"

OWNER = "gmp-e2e-owner@example.com"
DMS_MGR = "gmp-e2e-dms@example.com"
QA_MGR = "gmp-e2e-qa@example.com"
QC_MGR = "gmp-e2e-qc@example.com"
DEPT_MGR = "gmp-e2e-deptmgr@example.com"
EMP = "gmp-e2e-emp@example.com"
OUTSIDER = "gmp-e2e-outsider@example.com"


def _ensure_department(name, abbr):
    if not frappe.db.exists("Department", name):
        d = frappe.new_doc("Department")
        d.department_name = name
        d.is_group = 0
        d.flags.ignore_mandatory = True
        d.insert(ignore_permissions=True)
    frappe.db.set_value("Department", name, "custom_abbr", abbr)


def _ensure_type(code, label):
    if not frappe.db.exists("GMP Document Type", code):
        frappe.get_doc({"doctype": "GMP Document Type", "code": code, "type_name": label}).insert(
            ignore_permissions=True
        )


def _ensure_template(title, mappings):
    if frappe.db.exists("GMP Word Template", title):
        return title
    frappe.get_doc(
        {
            "doctype": "GMP Word Template",
            "template_title": title,
            "field_mappings": [{"custom_tag": t, "system_field": f} for t, f in mappings],
        }
    ).insert(ignore_permissions=True)
    return title


def _ensure_user(email, roles):
    if not frappe.db.exists("User", email):
        u = frappe.new_doc("User")
        u.email = email
        u.first_name = email.split("@")[0]
        u.user_type = "System User"
        u.send_welcome_email = 0
        u.insert(ignore_permissions=True)
    existing = set(
        frappe.get_all("Has Role", filters={"parent": email, "parenttype": "User"}, pluck="role")
    )
    for role in roles:
        if role not in existing:
            hr = frappe.new_doc("Has Role")
            hr.name = frappe.generate_hash(length=12)
            hr.update({"parent": email, "parenttype": "User", "parentfield": "roles", "role": role})
            hr.db_insert()
    frappe.db.commit()
    frappe.clear_cache(user=email)


def _ensure_employee(email, dept):
    name = frappe.db.get_value("Employee", {"user_id": email}, "name")
    if name:
        frappe.db.set_value("Employee", name, "department", dept)
        return name
    e = frappe.new_doc("Employee")
    e.name = f"GMP-E2E-EMP-{frappe.generate_hash(length=8)}"
    e.first_name = email.split("@")[0]
    e.employee_name = email.split("@")[0]
    e.user_id = email
    e.department = dept
    e.status = "Active"
    e.flags.ignore_mandatory = True
    e.db_insert()
    frappe.db.commit()
    return e.name


_SIG_PNG = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\xf8\x0f\x00"
    b"\x01\x01\x01\x00\x18\xdd\x8d\xb0\x00\x00\x00\x00IEND\xaeB`\x82"
)


def _ensure_signature(user, dept):
    """Give `user` a linked Employee + PNG signature so the Reviewer/QA
    signature validation passes for documents assigning this user."""
    emp = _ensure_employee(user, dept)
    if not frappe.db.get_value("Employee", emp, "custom_signature_image"):
        f = frappe.get_doc(
            {"doctype": "File", "file_name": f"sig-{frappe.generate_hash(length=6)}.png",
             "is_private": 1, "content": _SIG_PNG}
        ).insert(ignore_permissions=True)
        frappe.db.set_value("Employee", emp, "custom_signature_image", f.file_url)
    frappe.db.commit()


def _purge_docs():
    for name in frappe.get_all(
        "GMP Document", filters=[["document_name_en", "like", "GMP-E2E-%"]], pluck="name"
    ):
        try:
            frappe.db.set_value("GMP Document", name, "docstatus", 0, update_modified=False)
            frappe.delete_doc("GMP Document", name, ignore_permissions=True, force=True)
        except Exception:
            pass
    frappe.db.commit()


@unittest.skipUnless(SOFFICE_AVAILABLE, "LibreOffice (soffice) not on PATH")
class TestE2EPDF(FrappeTestCase):
    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        _ensure_department(DEPT_QA, "EQA")
        _ensure_department(DEPT_PROD, "EPR")
        _ensure_type("SOP", "SOP")
        _ensure_template(TPL_ALPHA, [("my_title", "document_name_en"), ("rev", "version_number")])
        _ensure_template(TPL_BETA, [("heading", "document_name_en"), ("eff", "effective_date")])
        for email, roles in (
            (OWNER, ["QA Manager"]),
            (DMS_MGR, ["DMS Manager"]),
            (QA_MGR, ["QA Manager"]),
            (QC_MGR, ["QC Manager"]) if frappe.db.exists("Role", "QC Manager") else (QC_MGR, ["Employee"]),
            (DEPT_MGR, ["Employee"]),
            (EMP, ["Employee"]),
            (OUTSIDER, []),
        ):
            _ensure_user(email, roles)
        _ensure_employee(EMP, DEPT_QA)
        _ensure_employee(QC_MGR, DEPT_QA)
        _ensure_employee(DEPT_MGR, DEPT_QA)
        # reviewer/qa_approver are Administrator in these tests; the signature
        # validation requires Administrator to have a signature image.
        _ensure_signature("Administrator", DEPT_QA)
        _purge_docs()

    @classmethod
    def tearDownClass(cls):
        _purge_docs()
        super().tearDownClass()

    def tearDown(self):
        frappe.set_user("Administrator")
        frappe.flags.pop("dms_user_departments", None)
        _purge_docs()

    # ----------------------------- helpers ----------------------------- #

    def _docx(self, fname, paragraphs):
        from docx import Document as DocxDocument
        import os
        import tempfile

        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, fname)
            d = DocxDocument()
            for p in paragraphs:
                d.add_paragraph(p)
            # Unique marker so two docx with the same visible text never produce
            # byte-identical files (Frappe dedups File records by content hash,
            # which would otherwise collapse a "fresh" revision onto the
            # predecessor's File and trip the amend inherited-attachment guard).
            d.add_paragraph(f"nonce:{frappe.generate_hash(length=20)}")
            d.save(path)
            with open(path, "rb") as fh:
                return frappe.get_doc(
                    {"doctype": "File", "file_name": fname, "is_private": 1, "content": fh.read()}
                ).insert(ignore_permissions=True)

    def _create_approved(self, en, template, paragraphs, dept=DEPT_QA, prepared_by=OWNER):
        docx = self._docx(f"{en}.docx", paragraphs)
        doc = frappe.new_doc("GMP Document")
        doc.update(
            {
                "document_name_fa": "تست",
                "document_name_en": en,
                "document_type": "SOP",
                "department": dept,
                "gmp_impact": "Major",
                "validity_period": "3 Years",
                "version_number": 0,
                "word_template": template,
                "reviewer": "Administrator",
                "qa_approver": "Administrator",
                "prepared_by": prepared_by,
            }
        )
        doc.attachment_file = docx.file_url
        doc.insert(ignore_permissions=True)
        return self._approve(doc)

    def _docx_exact(self, fname, paragraphs):
        """Like _docx but WITHOUT a uniqueness nonce, so two documents with the
        same paragraphs produce byte-identical files — exercising Frappe's
        content-hash deduplication (the real-world trigger for the content
        mixup, e.g. users starting each version from the same base file)."""
        from docx import Document as DocxDocument
        import os
        import tempfile

        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, fname)
            d = DocxDocument()
            for p in paragraphs:
                d.add_paragraph(p)
            d.save(path)
            with open(path, "rb") as fh:
                return frappe.get_doc(
                    {"doctype": "File", "file_name": fname, "is_private": 1, "content": fh.read()}
                ).insert(ignore_permissions=True)

    def _create_approved_exact(self, en, template, paragraphs, dept=DEPT_QA):
        docx = self._docx_exact(f"{en}.docx", paragraphs)
        doc = frappe.new_doc("GMP Document")
        doc.update(
            {
                "document_name_fa": "تست",
                "document_name_en": en,
                "document_type": "SOP",
                "department": dept,
                "gmp_impact": "Major",
                "validity_period": "3 Years",
                "version_number": 0,
                "word_template": template,
                "reviewer": "Administrator",
                "qa_approver": "Administrator",
                "prepared_by": OWNER,
            }
        )
        doc.attachment_file = docx.file_url
        doc.insert(ignore_permissions=True)
        return self._approve(doc)

    def test_identical_uploads_render_independently(self):
        """Two documents whose uploaded .docx are byte-identical (so Frappe
        deduplicates them onto one physical file) must each render their OWN
        content. Regression for the content mixup where the clean-render
        overwrite left File.content_hash stale, poisoning dedup so a later
        upload was pointed at an already-rendered file and one document's PDF
        contained another's content."""
        body = ["CONTROLLED_BODY_TEXT", "Title: {{ my_title }}"]
        a = self._create_approved_exact("GMP-E2E-Dup-A", TPL_ALPHA, body)
        b = self._create_approved_exact("GMP-E2E-Dup-B", TPL_ALPHA, body)

        ta = self._pdf_text(a.name)
        tb = self._pdf_text(b.name)
        # Each PDF must carry its own mapped title (my_title -> document_name_en).
        self.assertIn("GMP-E2E-Dup-A", ta, "A's PDF lost its own content")
        self.assertIn("GMP-E2E-Dup-B", tb, "B's PDF lost its own content")
        # And must NOT carry the other document's content.
        self.assertNotIn("GMP-E2E-Dup-B", ta, "A's PDF leaked B's content")
        self.assertNotIn("GMP-E2E-Dup-A", tb, "B's PDF leaked A's content (dedup mixup)")

    def test_amend_with_derived_upload_renders_new_version(self):
        """User's reported scenario: amend an approved document by re-uploading
        a .docx identical to the original (e.g. started from the same base
        file, which Frappe deduplicates). The new version's PDF must render the
        NEW version's data, not the predecessor's content."""
        body = ["AMEND_BODY_TEXT", "Version: {{ rev }}"]  # rev -> version_number
        v0 = self._create_approved_exact("GMP-E2E-AmendDup", TPL_ALPHA, body)
        self.assertIn("Version: 0", self._pdf_text(v0.name))

        v0.cancel()
        amended = frappe.copy_doc(v0)
        amended.docstatus = 0
        amended.amended_from = v0.name
        amended.reason_for_change = "Revision started from the same base file"
        # Byte-identical to v0's original upload -> exercises dedup.
        amended.attachment_file = self._docx_exact("GMP-E2E-AmendDup-v1.docx", body).file_url
        amended.insert(ignore_permissions=True)
        amended = self._approve(amended)

        text = self._pdf_text(amended.name)
        self.assertIn("Version: 1", text, "amended PDF did not render the new version")
        self.assertNotIn("Version: 0", text, "amended PDF leaked the predecessor's content")
        # The predecessor's controlled file must be intact and independent.
        self.assertIn("Version: 0", self._pdf_text(v0.name), "predecessor PDF was corrupted")

    def _add_reference(self, parent_name, target_name):
        frappe.get_doc(
            {
                "doctype": "GMP Document Reference",
                "parenttype": "GMP Document",
                "parentfield": "references",
                "parent": parent_name,
                "referenced_document": target_name,
                "reference_type": "References",
            }
        ).insert(ignore_permissions=True)

    def _approve(self, doc):
        apply_workflow(doc, "Submit for Review")
        apply_workflow(doc, "Approve as Reviewer")
        apply_workflow(doc, "Approve as QA")
        doc.reload()
        return doc

    def _base_pdf_path(self, docname):
        f = frappe.db.get_value(
            "File",
            {"attached_to_doctype": "GMP Document", "attached_to_name": docname, "file_name": f"{docname}.pdf"},
            "name",
        )
        self.assertTrue(f, f"base PDF File missing for {docname}")
        return frappe.get_doc("File", f).get_full_path()

    def _pdf_text(self, docname):
        from pypdf import PdfReader

        reader = PdfReader(self._base_pdf_path(docname))
        return "\n".join((p.extract_text() or "") for p in reader.pages)

    def _pdf_bytes(self, docname):
        with open(self._base_pdf_path(docname), "rb") as fh:
            return fh.read()

    # --------------------------- 8A: templates ------------------------- #

    def test_8a_template_differentiation_no_cross_contamination(self):
        a = self._create_approved(
            "GMP-E2E-Alpha-Doc",
            TPL_ALPHA,
            ["ALPHA_UNIQUE_HEADER", "Title: {{ my_title }}", "Rev: {{ rev }}"],
        )
        b = self._create_approved(
            "GMP-E2E-Beta-Doc",
            TPL_BETA,
            ["BETA_UNIQUE_FOOTER", "Heading: {{ heading }}", "Eff: {{ eff }}"],
        )

        ta, tb = self._pdf_text(a.name), self._pdf_text(b.name)
        # Distinct static content rendered into each PDF.
        self.assertIn("ALPHA_UNIQUE_HEADER", ta)
        self.assertIn("BETA_UNIQUE_FOOTER", tb)
        # No cross-contamination between templates/docs.
        self.assertNotIn("BETA_UNIQUE_FOOTER", ta)
        self.assertNotIn("ALPHA_UNIQUE_HEADER", tb)
        # Template-specific tag mapping was applied (my_title/heading ->
        # document_name_en), so each PDF carries its own mapped title.
        self.assertIn(a.document_name_en, ta)
        self.assertIn(b.document_name_en, tb)
        # The generated PDFs are not byte-identical.
        self.assertNotEqual(self._pdf_bytes(a.name), self._pdf_bytes(b.name))

    # --------------------------- 8B: versions -------------------------- #

    def test_8b_multi_version_content_correctness(self):
        texts = {}
        en = "GMP-E2E-Versioned"
        doc = self._create_approved(
            en, TPL_ALPHA, ["VERSIONED_DOC", "Version: {{ rev }}", "ID: {{ docname }}"]
        )
        texts[doc.version_number] = self._pdf_text(doc.name)

        # Revise four times: cancel + amend + re-approve, each a new version.
        for _ in range(4):
            doc.cancel()
            amended = frappe.copy_doc(doc)
            amended.docstatus = 0
            amended.amended_from = doc.name
            amended.reason_for_change = f"Revision to v{(doc.version_number or 0) + 1}"
            amended.attachment_file = self._docx(
                f"{en}-v{(doc.version_number or 0) + 1}.docx",
                ["VERSIONED_DOC", "Version: {{ rev }}", "ID: {{ docname }}"],
            ).file_url
            amended.insert(ignore_permissions=True)
            doc = self._approve(amended)
            texts[doc.version_number] = self._pdf_text(doc.name)

        # Five versions: 0..4.
        self.assertEqual(sorted(texts), [0, 1, 2, 3, 4])
        # Each PDF carries its own version number, and never a stale one.
        for v, text in texts.items():
            self.assertIn(f"Version: {v}", text, f"v{v} PDF missing its version number")
            for other in texts:
                if other != v:
                    self.assertNotIn(f"Version: {other}", text, f"v{v} PDF leaked v{other}")
        # All five PDFs differ pairwise.
        names = sorted(texts)
        for i in range(len(names)):
            for j in range(i + 1, len(names)):
                self.assertNotEqual(texts[names[i]], texts[names[j]])

    # ------------------------ 8C: independence ------------------------- #

    def test_8c_independent_documents_unique_and_isolated(self):
        d1 = self._create_approved("GMP-E2E-Indep-1", TPL_ALPHA, ["DOC_ONE_BODY", "{{ my_title }}"])
        d2 = self._create_approved("GMP-E2E-Indep-2", TPL_ALPHA, ["DOC_TWO_BODY", "{{ my_title }}"])
        d3 = self._create_approved("GMP-E2E-Indep-3", TPL_BETA, ["DOC_THREE_BODY", "{{ heading }}"])

        bytes_ = {n: self._pdf_bytes(n) for n in (d1.name, d2.name, d3.name)}
        # All unique.
        self.assertEqual(len({v for v in bytes_.values()}), 3)
        # Reference isolation: only d1 references d2; trees must reflect that.
        # d1 is submitted, so insert the reference child row directly rather
        # than re-saving the submitted document.
        self._add_reference(d1.name, d2.name)
        t1 = get_document_reference_tree(d1.name)
        t3 = get_document_reference_tree(d3.name)
        self.assertIn(d2.name, [c["name"] for c in t1["children"]])
        self.assertEqual(t3["children"], [])

    # --------------------------- 9: role access ------------------------ #

    def test_9_role_based_access_to_real_pdf(self):
        doc = self._create_approved("GMP-E2E-RoleDoc", TPL_ALPHA, ["ROLE_DOC", "{{ my_title }}"], dept=DEPT_QA)

        def can_download(user):
            frappe.set_user(user)
            try:
                download_watermarked_pdf(doc.name)
                return True
            except frappe.PermissionError:
                return False
            finally:
                frappe.set_user("Administrator")

        def can_read(user):
            frappe.set_user(user)
            try:
                return frappe.has_permission("GMP Document", "read", doc=doc.name)
            finally:
                frappe.set_user("Administrator")

        # DMS Manager: full access incl. PDF.
        self.assertTrue(can_read(DMS_MGR))
        self.assertTrue(can_download(DMS_MGR))
        # QA Manager: full access.
        self.assertTrue(can_download(QA_MGR))
        # Employee in the document's department: read + controlled-copy PDF.
        self.assertTrue(can_read(EMP))
        self.assertTrue(can_download(EMP))
        # Outsider (no role, no department): denied.
        self.assertFalse(can_read(OUTSIDER))
        self.assertFalse(can_download(OUTSIDER))

    def test_9_employee_write_denied_clean_word_manager_only(self):
        doc = self._create_approved("GMP-E2E-WordDoc", TPL_ALPHA, ["WORD_DOC", "{{ my_title }}"])
        # Employee may read but not edit.
        frappe.set_user(EMP)
        try:
            self.assertFalse(frappe.has_permission("GMP Document", "write", doc=doc.name))
            with self.assertRaises(frappe.PermissionError):
                download_word_document(doc.name)  # clean Word is manager-only
        finally:
            frappe.set_user("Administrator")
        # DMS Manager may download the clean Word.
        frappe.set_user(DMS_MGR)
        try:
            download_word_document(doc.name)
            self.assertEqual(frappe.local.response.type, "download")
        finally:
            frappe.set_user("Administrator")

    def test_9_cross_department_member_denied(self):
        prod = self._create_approved(
            "GMP-E2E-ProdDoc", TPL_ALPHA, ["PROD_DOC", "{{ my_title }}"], dept=DEPT_PROD
        )
        frappe.set_user(EMP)  # EMP is in QA dept, not PROD
        try:
            self.assertFalse(frappe.has_permission("GMP Document", "read", doc=prod.name))
            with self.assertRaises(frappe.PermissionError):
                download_watermarked_pdf(prod.name)
            with self.assertRaises(frappe.PermissionError):
                get_document_reference_tree(prod.name)
        finally:
            frappe.set_user("Administrator")

    # --------------------------- 10: regressions ----------------------- #

    def test_10_amend_available_after_cancel_for_qa_manager(self):
        """Regression (v1.2.5): after an approved document is cancelled, a plain
        QA Manager must be able to create a new version. The Amend button is
        gated by Frappe's can_amend(): docstatus == 2, amend permission, and the
        form is NOT workflow-read-only (the user shares a role with the current
        state's allow_edit). Cancelling puts the doc in the Obsolete state, so a
        QA Manager who lacks Obsolete's allow_edit role would see no Amend
        button. This verifies the full gate end-to-end and then performs the
        amend to confirm a new version is produced."""
        doc = self._create_approved(
            "GMP-E2E-AmendAfterCancel", TPL_ALPHA, ["AMEND_DOC", "{{ my_title }}"]
        )
        doc.cancel()
        doc.reload()
        self.assertEqual(doc.docstatus, 2)
        self.assertEqual(doc.workflow_status, "Obsolete")

        # 1) Amend-button visibility gate — mirror frappe.workflow.is_read_only:
        #    the form (and the Amend action) is read-only unless the user shares
        #    a role with the current workflow state's allow_edit roles.
        obsolete_roles = [
            s.allow_edit
            for s in frappe.get_doc("Workflow", "GMP Document Workflow").states
            if s.state == "Obsolete"
        ]
        qa_roles = set(frappe.get_roles(QA_MGR))
        workflow_read_only = not qa_roles.intersection(obsolete_roles)
        self.assertFalse(
            workflow_read_only,
            "QA Manager is workflow-read-only on a cancelled doc -> Amend button hidden",
        )
        # 2) amend permission (the other half of can_amend()).
        self.assertTrue(
            frappe.has_permission("GMP Document", "amend", doc=doc.name, user=QA_MGR),
            "QA Manager lacks amend permission",
        )

        # 3) The amend (what the button triggers) produces a new draft version.
        amended = frappe.copy_doc(doc)
        amended.docstatus = 0
        amended.amended_from = doc.name
        amended.reason_for_change = "Revised after cancellation"
        amended.attachment_file = self._docx(
            "GMP-E2E-AmendAfterCancel-v1.docx", ["AMEND_DOC_V1", "{{ my_title }}"]
        ).file_url
        amended.insert(ignore_permissions=True)
        self.assertEqual(amended.version_number, (doc.version_number or 0) + 1)
        self.assertEqual(amended.docstatus, 0)
        self.assertEqual(amended.amended_from, doc.name)

    def test_10_direct_submit_bypass_blocked(self):
        # A QA Manager must not be able to skip the workflow by calling submit()
        # directly; before_submit enforces workflow_status == Approved.
        docx = self._docx("GMP-E2E-Bypass.docx", ["BYPASS", "{{ my_title }}"])
        doc = frappe.new_doc("GMP Document")
        doc.update(
            {
                "document_name_fa": "ت",
                "document_name_en": "GMP-E2E-Bypass",
                "document_type": "SOP",
                "department": DEPT_QA,
                "gmp_impact": "Minor",
                "validity_period": "2 Years",
                "version_number": 0,
                "word_template": TPL_ALPHA,
                "reviewer": "Administrator",
                "qa_approver": "Administrator",
            }
        )
        doc.attachment_file = docx.file_url
        doc.insert(ignore_permissions=True)
        with self.assertRaises(frappe.PermissionError):
            doc.submit()

    def test_10_signature_inserted_changes_pdf(self):
        # Administrator (the approver in these tests) has a signature image
        # (ensured in setUpClass), so the signed base PDF embeds it as an image.
        from pypdf import PdfReader

        signed = self._create_approved(
            "GMP-E2E-Signed", TPL_ALPHA, ["SIGNED_DOC", "Approved by: {{ qa_signature }}"]
        )
        path = self._base_pdf_path(signed.name)
        imgs = sum(len(getattr(p, "images", []) or []) for p in PdfReader(path).pages)
        self.assertGreaterEqual(imgs, 1, "approver signature image was not embedded in the PDF")
